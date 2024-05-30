import openpyxl
from openai import OpenAI
from docx import Document
import os

# 请替换为你的OpenAI API密钥
client = OpenAI(
    base_url='https://api.openai-proxy.live/v1',
    api_key='sk-QNLQxq0rkkZcba18VLM80z1eguzLr3kbGua0WWsqpHjGl13E',
)

def generate_text_with_gpt3(title, prompt_words):
    """
    使用GPT-3根据标题和提示词生成文章
    """
    response = OpenAI.Completion.create(
      engine="gpt-3.5-turbo-0125", # 根据最新的API版本，这里可能需要更新
      prompt=f"{prompt_words}\n\nTitle: {title}\n\nArticle:",
      temperature=0.7,
      max_tokens=10000,
      top_p=1.0,
      frequency_penalty=0.0,
      presence_penalty=0.0
    )
    return response.choices[0].text.strip()

def read_titles_from_excel(excel_path):
    """
    从Excel文件读取标题
    """
    workbook = openpyxl.load_workbook(excel_path)
    sheet = workbook.active
    titles = [cell.value for row in sheet.iter_rows(min_row=2, values_only=True) for cell in row if cell]
    return titles

def create_and_save_document(title, content, folder="generated_articles"):
    """
    创建并保存文章到Word文档
    """
    os.makedirs(folder, exist_ok=True)
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    file_path = os.path.join(folder, f"{title[:30]}.docx".replace("/", "-").replace("\\", "-")) # 避免文件名非法字符
    doc.save(file_path)

def main():
    excel_path = 'path/to/your/excel.xlsx' # 更新为你的Excel文件路径
    prompt_words = "这里填入你的提示词" # 根据需要更新提示词
    titles = read_titles_from_excel(excel_path)

    for title in titles:
        content = generate_text_with_gpt3(title, prompt_words)
        create_and_save_document(title, content)

if __name__ == "__main__":
    main()
