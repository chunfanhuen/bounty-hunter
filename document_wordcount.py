import os
from docx import Document


def count_chinese_characters(text):
    count = 0
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            count += 1
    return count


def count_words_in_paragraph(paragraph):
    return count_chinese_characters(paragraph.text)


def count_words_in_table(table):
    word_count = 0
    for row in table.rows:
        for cell in row.cells:
            word_count += count_words_in_paragraph(cell)
    return word_count


def count_words_in_docx(docx_path):
    doc = Document(docx_path)
    word_count = 0

    for para in doc.paragraphs:
        word_count += count_words_in_paragraph(para)

    for table in doc.tables:
        word_count += count_words_in_table(table)

    return word_count


def find_files_with_less_than_100_words(folder_path):
    files_with_few_words = []

    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(folder_path, filename)
            word_count = count_words_in_docx(file_path)
            if word_count < 100:
                files_with_few_words.append(filename)

    return files_with_few_words


# 示例使用
folder_path = r'C:\Users\SFZDSB\Desktop\output'
files_with_few_words = find_files_with_less_than_100_words(folder_path)

if files_with_few_words:
    print('Files with less than 100 words:')
    for file in files_with_few_words:
        print(file)
else:
    print('No files with less than 100 words found.')
