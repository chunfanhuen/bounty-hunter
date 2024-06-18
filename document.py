import os
from docx import Document

"""
这个程序只是单纯在指定文件夹路径里面生成名字从一到六十的docx文件
方便测试用
"""

# 指定你想要生成文件的文件夹路径
folder_path = r"C:\Users\14534\Documents\num_file"

# 检查文件夹是否存在，如果不存在则创建它
if not os.path.exists(folder_path):
    os.makedirs(folder_path)

# 生成60个docx文件，文件名从1到60
for i in range(1, 121):
    doc = Document()
    doc.add_heading(f'Document {i}', 0)
    doc.add_paragraph(f'This is the content of document {i}.')

    file_path = os.path.join(folder_path, f'{i}.docx')
    doc.save(file_path)

print(f"{i} 个 docx 文件已生成在指定文件夹。")
