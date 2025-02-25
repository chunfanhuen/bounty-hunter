import os

import openai
import pandas as pd
from docx import Document

# 设置OpenAI的API密钥和base_url
client = openai.OpenAI(
    base_url='https://api.openai-proxy.live/v1',
    api_key='sk-QNLQxq0rkkZcba18VLM80z1eguzLr3kbGua0WWsqpHjGl13E',
)

# 提示词
prompt_file = r'C:\Users\14534\Desktop\prompt.txt'  # 提示词路径
with open(prompt_file, 'r', encoding='utf-8') as file:
    prompt_prefix = file.read()

# 读取Excel文件
excel_file = r'C:\Users\14534\Desktop\title.xlsx'
df = pd.read_excel(excel_file,usecols='A',header=None)
# 指定保存文件的文件夹路径
output_folder = r'C:\Users\14534\Desktop\output'
error_titles = []
# 循环读取每一个标题，并生成文章
for index, row in df.iterrows():
    title = row.iloc[0]  # 获取标题列
    prompt = f"{prompt_prefix}{title}"
    try:
        # 调用OpenAI API生成文章
        response = chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                    "max_tokens": "3000",
                }
            ],
            model="gpt-3.5-turbo",
        )
        print(response)
        article_content = response.choices[0].message.content.strip()

        # 创建docx文件并保存内容
        doc = Document()
        # doc.add_heading(title, 0)  # 文章标题可删除
        doc.add_paragraph(article_content)

        # 文件命名为“数字+题目”
        file_name = f"{index + 1}_{title}.docx"
        file_path = os.path.join(output_folder, file_name)
        doc.save(file_path)

        print(f"Article '{title}' has been saved as '{file_name}'.")
    except Exception as e:
        print(f"无法生成文档的标题名是：'{title}':{e}")
        error_titles.append(title)
print("All articles have been generated and saved.")

if error_titles:
    print("无法保存的标题是：")
    for error_title in error_titles:
        print(error_title)